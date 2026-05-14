from io import BytesIO
from typing import Dict

from docx import Document

from core.cdp import (
    disciplina_da_linha,
    habilidade_item_cdp,
    montar_acessibilidade_cdp,
    montar_acompanhamento_cdp,
    montar_metodologia_cdp,
    selecionar_item,
    titulo_item_cdp,
)
from docx_generator.preencher import (
    _eh_cabecalho_plano,
    _preencher_cabecalho,
    _preencher_celula_aprendizagem,
    _preencher_celula_lista,
    _preencher_celula_metodologia,
)


def _linha_de_aula_cdp(row) -> bool:
    if len(row.cells) < 5:
        return False
    primeira = (row.cells[0].text or "").strip().lower()
    material = (row.cells[1].text or "").strip().lower()
    if not primeira or not material:
        return False
    if "aula" in primeira and "semanal" in primeira:
        return False
    if "número e título" in material or "numero e titulo" in material:
        return False
    return bool(disciplina_da_linha(material))


def _first_distinct_cell_index(row, start_idx: int):
    if start_idx >= len(row.cells):
        return None
    vistos = {row.cells[i]._tc for i in range(start_idx + 1)}
    for idx in range(start_idx + 1, len(row.cells)):
        if row.cells[idx]._tc not in vistos:
            return idx
    return None


def _indices_cdp(row) -> Dict[str, int | None]:
    desenvolvimento_idx = 3
    acompanhamento_idx = _first_distinct_cell_index(row, desenvolvimento_idx)
    acessibilidade_idx = None
    if acompanhamento_idx is not None:
        vistos = {row.cells[i]._tc for i in range(acompanhamento_idx + 1)}
        for idx in range(len(row.cells) - 1, acompanhamento_idx, -1):
            if row.cells[idx]._tc not in vistos:
                acessibilidade_idx = idx
                break
    return {
        "material": 1,
        "aprendizagem": 2,
        "desenvolvimento": desenvolvimento_idx,
        "acompanhamento": acompanhamento_idx,
        "acessibilidade": acessibilidade_idx,
    }


def _disciplina_exibicao(disciplina: str) -> str:
    nomes = {
        "português": "PORTUGUÊS",
        "matematica": "MATEMÁTICA",
        "história": "HISTÓRIA",
        "geografia": "GEOGRAFIA",
        "ciências": "CIÊNCIAS",
        "arte": "ARTE",
    }
    return nomes.get(disciplina, disciplina.upper())


def _habilidade(item: Dict[str, str]) -> str:
    habilidade = habilidade_item_cdp(item)
    return f"HABILIDADE:\n{habilidade}" if habilidade else ""


def _material(disciplina: str, item: Dict[str, str]) -> str:
    titulo = titulo_item_cdp(item)
    return f"{_disciplina_exibicao(disciplina)}\nTema: {titulo}"


def _metodologia_dict(texto: str):
    return [{"titulo": "", "texto": texto}]


def _preencher_cabecalhos_cdp(
    doc,
    professor: str,
    componente: str,
    turma: str,
    mes: str,
    bimestre: str,
    observacao: str,
    aulas_previstas_manual: str,
) -> None:
    for table in doc.tables:
        if _eh_cabecalho_plano(table):
            _preencher_cabecalho(
                table,
                professor,
                componente,
                turma,
                mes,
                bimestre,
                observacao,
                aulas_previstas_manual,
            )


def preencher_documento_cdp(
    modelo_docx,
    professor: str,
    turma: str,
    mes: str = "",
    bimestre: str = "",
    aula_inicial: int = 1,
    fundamental: bool = False,
    multisseriada: bool = False,
    serie_cdp: str = "",
    componente_cdp: str = "",
    item_cdp: Dict[str, str] | None = None,
    observacao: str = "",
    aulas_previstas_manual: str = "",
) -> BytesIO:
    doc = Document(modelo_docx)
    componente = "CDP - ENSINO FUNDAMENTAL" if fundamental else "MULTISSERIADA - EJA FUNDAMENTAL - ANOS INICIAIS"
    _preencher_cabecalhos_cdp(
        doc,
        professor,
        componente,
        turma,
        mes,
        bimestre,
        observacao,
        aulas_previstas_manual,
    )

    contadores: Dict[str, int] = {}
    for table in doc.tables:
        for row in table.rows:
            if not _linha_de_aula_cdp(row):
                continue

            idxs = _indices_cdp(row)
            disciplina = disciplina_da_linha(row.cells[idxs["material"]].text)

            if item_cdp and contadores.get(disciplina, 0) == 0:
                item = item_cdp
            else:
                contador = contadores.get(disciplina, 0)
                item = selecionar_item(
                    disciplina,
                    contador,
                    turma=serie_cdp if multisseriada and serie_cdp else turma,
                    bimestre=bimestre,
                    aula_inicial=aula_inicial,
                    fundamental=fundamental,
                    multisseriada=multisseriada,
                    componente_cdp=componente_cdp,
                )

            contadores[disciplina] = contadores.get(disciplina, 0) + 1
            if not item:
                continue

            _preencher_celula_aprendizagem(row.cells[idxs["aprendizagem"]], _habilidade(item))
            _preencher_celula_metodologia(
                row.cells[idxs["desenvolvimento"]],
                _metodologia_dict(montar_metodologia_cdp(disciplina, item, fundamental=fundamental)),
            )
            row.cells[idxs["material"]].text = _material(disciplina, item)

            if idxs["acompanhamento"] is not None:
                _preencher_celula_lista(
                    row.cells[idxs["acompanhamento"]],
                    montar_acompanhamento_cdp(disciplina, item, fundamental=fundamental),
                )
            if idxs["acessibilidade"] is not None:
                _preencher_celula_lista(
                    row.cells[idxs["acessibilidade"]],
                    montar_acessibilidade_cdp(disciplina, item, fundamental=fundamental),
                )

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

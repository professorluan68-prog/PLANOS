from io import BytesIO
from datetime import date
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from core.helpers import texto_lista


DESTAQUES_TEXTO = {
    "VIREM E CONVERSEM": "Virem e conversem",
    "TODO MUNDO ESCREVE": "Todo mundo escreve",
    "COM SUAS PALAVRAS": "Com suas palavras",
    "HORA DA LEITURA": "Hora da leitura",
    "DE OLHO NO MODELO": "De olho no modelo",
    "UM PASSO DE CADA VEZ": "Um passo de cada vez",
    "LISTEN AND REPEAT": "Listen and repeat",
    "WRITE AND SHARE": "Write and share",
    "SAY IT IN ENGLISH": "Say it in English",
}

TITULOS_ETAPAS = {
    "Para comecar": "Para come\u00e7ar",
    "Contextualizacao": "Contextualiza\u00e7\u00e3o",
    "Leitura analitica": "Leitura anal\u00edtica",
    "Leitura e construcao do conteudo": "Leitura e constru\u00e7\u00e3o do conte\u00fado",
    "Exploracao": "Explora\u00e7\u00e3o",
    "Foco no conteudo": "Foco no conte\u00fado",
    "Formalizacao": "Formaliza\u00e7\u00e3o",
    "Sistematizacao": "Sistematiza\u00e7\u00e3o",
    "Sistematiza\u00e7\u00e3o": "Sistematiza\u00e7\u00e3o",
    "Pause e responda": "Pause e responda",
    "Na pratica": "Na pr\u00e1tica",
    "Analise de caso": "An\u00e1lise de caso",
    "Calculos financeiros": "C\u00e1lculos financeiros",
    "Planejamento orcamentario": "Planejamento or\u00e7ament\u00e1rio",
    "Projeto empreendedor": "Projeto empreendedor",
    "Revisao e reescrita": "Revis\u00e3o e reescrita",
    "Relembre": "Relembre",
    "Encerramento": "Encerramento",
    "Para come\u00e7ar": "Para come\u00e7ar",
    "Contextualiza\u00e7\u00e3o": "Contextualiza\u00e7\u00e3o",
    "Leitura anal\u00edtica": "Leitura anal\u00edtica",
    "Leitura e constru\u00e7\u00e3o do conte\u00fado": "Leitura e constru\u00e7\u00e3o do conte\u00fado",
    "Explora\u00e7\u00e3o": "Explora\u00e7\u00e3o",
    "Na pr\u00e1tica": "Na pr\u00e1tica",
    "An\u00e1lise de caso": "An\u00e1lise de caso",
    "C\u00e1lculos financeiros": "C\u00e1lculos financeiros",
    "Planejamento or\u00e7ament\u00e1rio": "Planejamento or\u00e7ament\u00e1rio",
    "Revis\u00e3o e reescrita": "Revis\u00e3o e reescrita",
}


def _substituir_texto(paragraph, substituicoes: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    texto_original = paragraph.text
    texto_novo = texto_original
    for chave, valor in substituicoes.items():
        texto_novo = texto_novo.replace(chave, valor)
    if texto_novo == texto_original:
        return
    paragraph.clear()
    paragraph.add_run(texto_novo)


def _substituir_em_tabela(tabela, substituicoes: dict[str, str]) -> None:
    for linha in tabela.rows:
        for celula in linha.cells:
            for paragrafo in celula.paragraphs:
                _substituir_texto(paragrafo, substituicoes)
            for tabela_interna in celula.tables:
                _substituir_em_tabela(tabela_interna, substituicoes)


def _texto_metodologia(aula: dict) -> str:
    metodologia = aula.get("metodologia") or []
    blocos = []
    for item in metodologia:
        if isinstance(item, dict):
            titulo = item.get("titulo", "")
            texto = item.get("texto", "")
            blocos.append(f"{titulo}\n{texto}".strip())
        else:
            blocos.append(str(item))
    return "\n\n".join(blocos)


# ── Constantes de formatação ────────────────────────────────────────────────
_FONTE_PADRAO = "Arial"
_TAMANHO_PADRAO = Pt(10)
_COR_VERMELHA = RGBColor(0xEE, 0x00, 0x00)
_PADRAO_BNCC = re.compile(r'(\([A-Z]{2}\d{2}[A-Z]{2,4}\d{0,3}[A-Z]?\))')
_PADRAO_TURMA_METODOLOGIA = re.compile(
    r"\b(da turma|com a turma)\s+\d{1,2}\s*[º°oªa?]?\s*(?:ano|s[ée]rie|em|ef)?\s*[A-Z]?\b",
    flags=re.I,
)

_CORRECOES_TEXTO_FINAL = {
    "participacao": "participação",
    "discussoes": "discussões",
    "discussao": "discussão",
    "analises": "análises",
    "analise": "análise",
    "interpretacao": "interpretação",
    "argumentacao": "argumentação",
    "ampliacao": "ampliação",
    "estrategias": "estratégias",
    "estrategia": "estratégia",
    "producao": "produção",
    "producoes": "produções",
    "vocabulario": "vocabulário",
    "relacoes": "relações",
    "necessarias": "necessárias",
    "necessarios": "necessários",
    "criterios": "critérios",
    "organizacao": "organização",
    "construcao": "constru\u00e7\u00e3o",
    "construcoes": "constru\u00e7\u00f5es",
    "correcao": "corre\u00e7\u00e3o",
    "correcoes": "corre\u00e7\u00f5es",
    "questao": "quest\u00e3o",
    "questoes": "quest\u00f5es",
    "solucao": "solu\u00e7\u00e3o",
    "solucoes": "solu\u00e7\u00f5es",
    "explicacao": "explica\u00e7\u00e3o",
    "explicacoes": "explica\u00e7\u00f5es",
    "pratica": "pr\u00e1tica",
    "praticas": "pr\u00e1ticas",
    "conteudos": "conte\u00fados",
    "mediacao": "mediação",
    "mediacoes": "mediações",
    "flexibilizacao": "flexibilização",
    "linguisticos": "linguísticos",
    "linguisticas": "linguísticas",
    "situacoes": "situações",
    "compreensao": "compreensão",
    "conteudo": "conteúdo",
    "genero": "gênero",
    "publico": "público",
    "previos": "prévios",
    "hipoteses": "hipóteses",
    "duvidas": "dúvidas",
    "sintese": "síntese",
    "sinteses": "sínteses",
    "importancia": "importância",
    "situacao": "situação",
    "equacao": "equação",
    "equacoes": "equações",
    "resolucao": "resolução",
    "resolucoes": "resoluções",
    "verificacao": "verificação",
    "conferencia": "conferência",
    "coerencia": "coerência",
    "incognita": "incógnita",
    "grafico": "gráfico",
    "graficos": "gráficos",
    "matematica": "matemática",
    "matematicas": "matemáticas",
    "calculos": "cálculos",
    "calculo": "cálculo",
    "validacao": "validação",
    "informacoes": "informações",
    "interpretacoes": "interpretações",
    "conclusoes": "conclusões",
    "proporcao": "proporção",
    "proporcoes": "proporções",
    "representacao": "representação",
    "representacoes": "representações",
    "dependencia": "dependência",
    "relacao": "relação",
    "comparacao": "comparação",
    "comparacoes": "comparações",
    "aplicacao": "aplicação",
    "aplicacoes": "aplicações",
    "responsaveis": "responsáveis",
    "responsavel": "responsável",
    "orcamento": "orçamento",
    "orcamentos": "orçamentos",
    "decisao": "decisão",
    "decisoes": "decisões",
    "experiencias": "experiências",
    "simulacoes": "simulações",
    "credito": "crédito",
    "creditos": "créditos",
    "divida": "dívida",
    "dividas": "dívidas",
    "emprestimo": "empréstimo",
    "emprestimos": "empréstimos",
    "instituicao": "instituição",
    "instituicoes": "instituições",
    "movimentacao": "movimentação",
    "poupanca": "poupança",
    "emergencia": "emergência",
    "familiares": "familiares",
    "habitos": "hábitos",
    "julgamentos": "julgamentos",
    "necessario": "necessário",
    "necessaria": "necessária",
    "vocabulario": "vocabulário",
    "aplicacao": "aplicação",
    "orcamentario": "orçamentário",
    "orcamentaria": "orçamentária",
}


def _aplicar_fonte(run, nome=_FONTE_PADRAO, tamanho=_TAMANHO_PADRAO, bold=None, color=None):
    """Aplica formatação padrão a um run."""
    run.font.name = nome
    run.font.size = tamanho
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def _limpar_celula(celula) -> None:
    celula.text = ""


def _paragrafo_base(celula):
    if not celula.paragraphs:
        return celula.add_paragraph()
    return celula.paragraphs[0]


def _normalizar_destaques(texto: str) -> str:
    texto_final = str(texto or "")
    for original, exibicao in DESTAQUES_TEXTO.items():
        texto_final = texto_final.replace(original, exibicao)
    return texto_final


def _capitalizar_como(original: str, corrigido: str) -> str:
    if original.isupper():
        return corrigido.upper()
    if original[:1].isupper():
        return corrigido[:1].upper() + corrigido[1:]
    return corrigido


def _polir_texto_docx(texto: str) -> str:
    texto_final = _PADRAO_TURMA_METODOLOGIA.sub(lambda m: m.group(1), str(texto or ""))
    texto_final = "\n".join(
        re.sub(r"[ \t\r\f\v]+", " ", linha).strip()
        for linha in texto_final.splitlines()
    ).strip()
    texto_final = re.sub(r"\bde o conceito\b", "do conceito", texto_final, flags=re.I)
    texto_final = re.sub(r"\b1o\b", "1º", texto_final, flags=re.I)
    texto_final = re.sub(r"\b1\s*o\s+grau\b", "1º grau", texto_final, flags=re.I)
    for sem_acento, com_acento in _CORRECOES_TEXTO_FINAL.items():
        texto_final = re.sub(
            rf"\b{re.escape(sem_acento)}\b",
            lambda m, novo=com_acento: _capitalizar_como(m.group(0), novo),
            texto_final,
            flags=re.I,
        )
    return texto_final


def _titulo_exibicao(titulo: str) -> str:
    return TITULOS_ETAPAS.get(str(titulo or "").strip(), str(titulo or "").strip())


def _adicionar_texto_com_destaques(paragrafo, texto: str) -> None:
    restante = _normalizar_destaques(texto)
    if not restante:
        return

    padrao = "|".join(re.escape(valor) for valor in DESTAQUES_TEXTO.values())
    if not padrao:
        paragrafo.add_run(restante)
        return

    partes = re.split(f"({padrao})", restante)
    for parte in partes:
        if not parte:
            continue
        run = paragrafo.add_run(parte)
        if parte in DESTAQUES_TEXTO.values():
            run.bold = True


def _preencher_celula_centralizada(celula, texto: str, bold: bool = False, color=None) -> None:
    _limpar_celula(celula)
    paragrafo = _paragrafo_base(celula)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragrafo.add_run(_polir_texto_docx(str(texto or "")))
    _aplicar_fonte(run, bold=bold, color=color)


def _preencher_celula_lista(celula, itens) -> None:
    _limpar_celula(celula)
    itens_lista = list(itens or [])
    if not itens_lista:
        return

    primeiro = _paragrafo_base(celula)
    for indice, item in enumerate(itens_lista):
        paragrafo = primeiro if indice == 0 else celula.add_paragraph()
        paragrafo.alignment = None
        texto_item = re.sub(r"^\s*(?:\u2611|\u2713|\u2022|-)\s*", "", str(item or "").strip())
        texto_item = _polir_texto_docx(texto_item)
        run_check = paragrafo.add_run("\u2611 ")
        run_check.font.name = "Segoe UI Symbol"
        run_check.font.size = _TAMANHO_PADRAO
        _adicionar_texto_com_destaques_formatado(paragrafo, texto_item)


def _adicionar_texto_com_destaques_formatado(paragrafo, texto: str) -> None:
    """Igual a _adicionar_texto_com_destaques mas aplica Arial 10pt a cada run."""
    restante = _normalizar_destaques(texto)
    if not restante:
        return
    padrao = "|".join(re.escape(valor) for valor in DESTAQUES_TEXTO.values())
    if not padrao:
        _aplicar_fonte(paragrafo.add_run(restante))
        return
    partes = re.split(f"({padrao})", restante)
    for parte in partes:
        if not parte:
            continue
        run = paragrafo.add_run(parte)
        _aplicar_fonte(run, bold=True if parte in DESTAQUES_TEXTO.values() else None)


def _preencher_celula_aprendizagem(celula, texto: str) -> None:
    """Preenche a coluna Aprendizagem: código BNCC em vermelho+bold, resto bold preto, centralizado."""
    _limpar_celula(celula)
    paragrafo = _paragrafo_base(celula)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    texto = _polir_texto_docx(str(texto or "").strip())
    if not texto:
        return
    # Tenta separar o código BNCC do texto descritivo
    match = _PADRAO_BNCC.search(texto)
    if match:
        codigo = match.group(1)
        pos = match.start()
        antes = texto[:pos].strip()
        depois = texto[pos + len(codigo):].strip()
        if antes:
            _aplicar_fonte(paragrafo.add_run(antes + " "), bold=True)
        _aplicar_fonte(paragrafo.add_run(codigo + " "), bold=True, color=_COR_VERMELHA)
        if depois:
            _aplicar_fonte(paragrafo.add_run(depois), bold=True)
    else:
        _aplicar_fonte(paragrafo.add_run(texto), bold=True)


def _preencher_celula_metodologia(celula, metodologia) -> None:
    _limpar_celula(celula)
    itens = list(metodologia or [])
    if not itens:
        return

    primeiro = _paragrafo_base(celula)
    primeiro.text = ""
    paragrafo_atual = primeiro

    for indice, item in enumerate(itens):
        if isinstance(item, dict):
            titulo = str(item.get("titulo") or "").strip()
            texto = str(item.get("texto") or "").strip()
            if titulo:
                texto_item = f"{_titulo_exibicao(_normalizar_destaques(titulo))}: {texto}"
            else:
                texto_item = texto
        else:
            texto_item = str(item).strip()
        texto_item = _polir_texto_docx(texto_item)

        # Separar por quebras de linha para lidar com o texto vindo da UI
        linhas = [l.strip() for l in texto_item.split('\n') if l.strip()]
        for linha in linhas:
            if not linha:
                continue

            if paragrafo_atual.text:
                paragrafo_atual = celula.add_paragraph()

            # Procurar por um padrão "Titulo: texto" para colocar em negrito
            match = re.match(r'^([^:]{2,35}):\s*(.*)$', linha)
            if match:
                titulo_bold = match.group(1) + ":"
                resto_texto = " " + match.group(2)
                _aplicar_fonte(paragrafo_atual.add_run(titulo_bold), bold=True)
                _adicionar_texto_com_destaques_formatado(paragrafo_atual, resto_texto)
            else:
                _adicionar_texto_com_destaques_formatado(paragrafo_atual, linha)


def _texto_tabela(tabela) -> str:
    texto = " ".join(celula.text.upper() for linha in tabela.rows for celula in linha.cells)
    return " ".join(texto.split())


def _eh_cabecalho_plano(tabela) -> bool:
    texto = _texto_tabela(tabela)
    return len(tabela.rows) >= 4 and "PLANO DE AULAS" in texto and "PROFESSOR" in texto


def _eh_tabela_aulas(tabela) -> bool:
    texto = _texto_tabela(tabela)
    return len(tabela.rows) >= 2 and "AULA SEMANAL" in texto and "APRENDIZAGEM" in texto


def _celulas_unicas(linha):
    vistas = set()
    celulas = []
    for celula in linha.cells:
        chave = id(celula._tc)
        if chave not in vistas:
            vistas.add(chave)
            celulas.append(celula)
    return celulas


def _definir_texto(celula, texto: str) -> None:
    celula.text = str(texto or "")


def _limpar_linha(linha) -> None:
    for celula in _celulas_unicas(linha):
        _definir_texto(celula, "")


def _formatar_data_horario(aula: dict) -> str:
    data = str(aula.get("data") or "").strip()
    horario = str(aula.get("horario") or "").strip()
    partes = [parte.strip() for parte in horario.splitlines() if parte.strip()]

    if len(partes) >= 2:
        return "\n".join([data, partes[1], partes[0]]).strip()
    if horario:
        return "\n".join([data, horario]).strip()
    return data


def _data_ddmm(texto: str):
    partes = str(texto or "").strip().split("/")
    if len(partes) < 2:
        return None
    try:
        return date(2000, int(partes[1]), int(partes[0]))
    except ValueError:
        return None


def _intervalo_cabecalho(tabela):
    if len(tabela.rows) < 4 or len(tabela.rows[3].cells) < 2:
        return None
    texto = tabela.rows[3].cells[1].text
    if " a " not in texto:
        return None
    inicio, fim = texto.split(" a ", 1)
    inicio_data = _data_ddmm(inicio)
    fim_data = _data_ddmm(fim)
    if not inicio_data or not fim_data:
        return None
    return inicio_data, fim_data


def _aula_pertence_ao_intervalo(aula: dict, intervalo) -> bool:
    data_aula = _data_ddmm(aula.get("data"))
    if not data_aula or not intervalo:
        return False
    inicio, fim = intervalo
    return inicio <= data_aula <= fim


def _titulo_aula(aula: dict, numero: int) -> str:
    titulo = str(aula.get("titulo") or aula.get("material") or aula.get("tema") or "").strip()
    if not titulo:
        return f"AULA {numero}"
    if titulo.upper().startswith("AULA"):
        return titulo
    return f"AULA {numero} - {titulo}"


def _preencher_cabecalho(
    tabela,
    professor: str,
    disciplina: str,
    turma: str,
    mes: str,
    bimestre: str,
    observacao: str,
    aulas_previstas: str,
) -> None:
    if len(tabela.rows) < 4:
        return

    linha_dados = tabela.rows[2].cells
    if len(linha_dados) >= 9:
        _definir_texto(linha_dados[2], professor)
        _definir_texto(linha_dados[3], disciplina)
        _definir_texto(linha_dados[6], turma)
        _definir_texto(linha_dados[7], mes)
        _definir_texto(linha_dados[8], bimestre)

    linha_semana = tabela.rows[3].cells
    if len(linha_semana) >= 4:
        _definir_texto(linha_semana[3], aulas_previstas)
    if observacao and len(linha_semana) >= 6:
        _definir_texto(linha_semana[5], observacao)


def _preencher_linha_aula(linha, aula: dict, numero: int) -> None:
    celulas = linha.cells
    if len(celulas) < 6:
        return

    # Col 0: Data/Horário — vermelho, centralizado, Arial 10
    _preencher_celula_centralizada(celulas[0], _formatar_data_horario(aula), color=_COR_VERMELHA)
    # Col 1: Título — vermelho + bold, centralizado, Arial 10
    _preencher_celula_centralizada(celulas[1], _titulo_aula(aula, numero), bold=True, color=_COR_VERMELHA)
    # Col 2: Aprendizagem — código BNCC vermelho, texto bold preto, centralizado
    _preencher_celula_aprendizagem(celulas[2], aula.get("aprendizagem", ""))
    # Col 3: Metodologia — título bold, texto normal, Arial 10
    _preencher_celula_metodologia(celulas[3], aula.get("metodologia"))
    # Col 4: Acompanhamento — Arial 10
    _preencher_celula_lista(celulas[4], aula.get("acompanhamento"))
    # Col 5: Acessibilidade — Arial 10
    _preencher_celula_lista(celulas[5], aula.get("acessibilidade"))


def _preencher_tabelas_modelo(
    documento,
    aulas,
    professor: str,
    disciplina: str,
    turma: str,
    mes: str,
    bimestre: str,
    observacao: str,
    aulas_previstas_manual: str,
) -> bool:
    pares = []
    tabelas = list(documento.tables)
    for indice, tabela in enumerate(tabelas):
        if not _eh_cabecalho_plano(tabela):
            continue
        proxima = tabelas[indice + 1] if indice + 1 < len(tabelas) else None
        if proxima is not None and _eh_tabela_aulas(proxima):
            pares.append((tabela, proxima))

    if not pares:
        return False

    aulas = list(aulas or [])
    aulas_por_par = [[] for _ in pares]
    usadas = set()

    for aula_indice, aula in enumerate(aulas):
        for par_indice, (cabecalho, _) in enumerate(pares):
            if _aula_pertence_ao_intervalo(aula, _intervalo_cabecalho(cabecalho)):
                aulas_por_par[par_indice].append((aula_indice + 1, aula))
                usadas.add(aula_indice)
                break

    sobras = [(indice + 1, aula) for indice, aula in enumerate(aulas) if indice not in usadas]
    for par_indice, (_, tabela_aulas) in enumerate(pares):
        vagas = max(0, len(tabela_aulas.rows) - 1 - len(aulas_por_par[par_indice]))
        if vagas and sobras:
            aulas_por_par[par_indice].extend(sobras[:vagas])
            sobras = sobras[vagas:]

    for par_indice, (cabecalho, tabela_aulas) in enumerate(pares):
        linhas_conteudo = list(tabela_aulas.rows[1:])
        aulas_da_semana = aulas_por_par[par_indice][: len(linhas_conteudo)]
        aulas_previstas = str(aulas_previstas_manual or len([a for a in aulas_da_semana if a])).strip()

        _preencher_cabecalho(
            cabecalho,
            professor,
            disciplina,
            turma,
            mes,
            bimestre,
            observacao,
            aulas_previstas,
        )

        for linha in linhas_conteudo:
            _limpar_linha(linha)

        for linha, (numero, aula) in zip(linhas_conteudo, aulas_da_semana):
            _preencher_linha_aula(linha, aula, numero)

    return True


def preencher_documento(
    modelo_stream,
    aulas,
    professor: str,
    disciplina: str,
    turma: str,
    mes: str,
    bimestre: str,
    observacao: str = "",
    aulas_previstas_manual: str = "",
):
    documento = Document(modelo_stream)
    primeira_aula = (aulas or [{}])[0]
    substituicoes = {
        "{{PROFESSOR}}": professor or "",
        "{{DISCIPLINA}}": disciplina or "",
        "{{TURMA}}": turma or "",
        "{{MES}}": mes or "",
        "{{BIMESTRE}}": bimestre or "",
        "{{OBSERVACAO}}": observacao or "",
        "{{AULAS_PREVISTAS}}": str(aulas_previstas_manual or ""),
        "{{TEMA}}": primeira_aula.get("tema", ""),
        "{{DATA}}": primeira_aula.get("data", ""),
        "{{HORARIO}}": primeira_aula.get("horario", ""),
        "{{APRENDIZAGEM}}": primeira_aula.get("aprendizagem", ""),
        "{{METODOLOGIA}}": _texto_metodologia(primeira_aula),
        "{{ACOMPANHAMENTO}}": texto_lista(primeira_aula.get("acompanhamento")),
        "{{ACESSIBILIDADE}}": texto_lista(primeira_aula.get("acessibilidade")),
    }
    for paragrafo in documento.paragraphs:
        _substituir_texto(paragrafo, substituicoes)
    for tabela in documento.tables:
        _substituir_em_tabela(tabela, substituicoes)

    _preencher_tabelas_modelo(
        documento,
        aulas,
        professor,
        disciplina,
        turma,
        mes,
        bimestre,
        observacao,
        aulas_previstas_manual,
    )

    saida = BytesIO()
    documento.save(saida)
    saida.seek(0)
    return saida
